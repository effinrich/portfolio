#!/usr/bin/env python3
"""Build an ATS-friendly DOCX + PDF resume for Rich Tillman from one content model.

Output directory defaults to the script's own location (public/resume/); override
with the RESUME_OUTDIR env var. The PDF font directory defaults to the Liberation
fonts path on Debian/Ubuntu; override with RESUME_FONT_DIR. If the fonts are not
found, the PDF falls back to the built-in Helvetica (Arial-metric) font.
"""

import os

OUTDIR = os.environ.get("RESUME_OUTDIR", os.path.dirname(os.path.abspath(__file__)))
BASENAME = "Rich-Tillman-Resume"
# Arial is the metric match for Liberation Sans and is universally available in Word.
DOCX_FONT = "Arial"

NAME = "Rich Tillman"
CONTACT = ("Elizabethton, TN  |  richtillman@pm.me  |  843-834-0041  |  "
           "linkedin.com/in/effinrich  |  github.com/effinrich  |  richtillman.xyz")
SUMMARY = ("Principal Frontend Engineer (15 yrs) specializing in React/TypeScript design systems and "
           "Storybook-driven development. Founder of ForgeKit, an open-source Figma-to-React CLI + MCP "
           "suite (5,700+ npm installs). IC-to-Director experience across five 0-to-1 startups. Seeking "
           "Principal/Senior Frontend, Product Engineer, and design-systems/dev-tooling roles.")

# Each job: (title, company, dates, [bullets])
JOBS = [
    ("Senior Frontend Engineer, AI Training & Evaluations",
     "Mercor / micro1 / Handshake AI", "November 2025 – June 2026", [
        "Designed programming prompts and evaluation rubrics scoring AI output on advanced frontend tasks—React architecture, TypeScript patterns, Storybook-driven development—with measurable model-quality gains traced to prompt design.",
        "Built AI-driven developer interfaces for interactive coding environments (React 19, Next.js, TanStack Start, Chakra UI, Storybook 10+, Nx), accelerating prototyping and workflow efficiency.",
        "Prototyped AI-driven developer interfaces with Lovable before hardening them in React 19 / TypeScript, compressing the design-to-working-UI loop.",
        "Advised engineering teams on integrating AI-assisted workflows (Claude Code, Cursor, Lovable) into production, streamlining code reviews and shortening release cycles.",
        "Implemented schema for AI evaluation rubrics and React/TypeScript architecture scoring.",
        "Utilized debug traces in agentic systems for Claude and OpenAI prompt and rubric checks.",
     ]),
    ("Staff Frontend Engineer & Tech Lead → Engineering Director",
     "Redesign Health", "July 2022 – May 2024", [
        "Promoted twice to Engineering Director; managed direct reports and coordinated async delivery across three cross-functional teams (15+ members) via Jira and Slack while staying hands-on, improving alignment and keeping releases on schedule.",
        "Spearheaded a 50+ component React design system in Storybook + Chromatic with a Next.js docs site—cutting dev time 30% across a 10–15 engineer org.",
        "Engineered 20+ TypeScript data-visualization components with custom hooks and Zustand (40% less rendering overhead) and established a Chromatic visual-regression workflow catching UI regressions pre-merge.",
        "Built a proprietary onboarding methodology (Storybook + feature-based architecture + custom hooks) that trained backend engineers to production-level React in one month.",
        "Devised a fixture-based mock-API strategy that unblocked a high-stakes demo under a 2-day deadline with zero integration friction while the backend built real endpoints in parallel.",
     ]),
    ("Senior Frontend Engineer & Tech Lead",
     "Pineapple Corporation", "January 2022 – July 2022", [
        "Orchestrated Nx monorepo adoption across 8+ React/TypeScript applications—35% version-control efficiency gain via shared libraries and Nx affected commands.",
        "Engineered cross-platform architecture in TypeScript (Expo, Nx, NativeBase, React Native)—25% iOS/Android performance gain for 100K+ users.",
        "Implemented a Storybook + Chromatic component-driven workflow for 60+ React components, accelerating stakeholder design reviews 50%.",
        "Led Redux Toolkit adoption across monorepo projects, standardizing data flow.",
     ]),
    ("Founding Frontend Engineer",
     "PHC Global", "July 2021 – January 2022", [
        "Built the foundational Nx monorepo for a B2B fintech dashboard (Next.js, React, 30+ shared libraries, gRPC middleware)—40% better developer experience.",
        "Refined UX across 8+ fintech workflows with responsive design, lifting NPS satisfaction 35%.",
        "Cut infrastructure costs 30% and time-to-market by 8 weeks with a scalable GCP backend (gRPC, Helm, Kubernetes).",
     ]),
    ("Lead Frontend Engineer",
     "Freebird", "September 2016 – January 2021", [
        "Built one of the earliest production React UI libraries on Storybook—200 components across B2B, B2C, and React Native—eliminating separate iOS and Android teams.",
        "Acted as design-engineering liaison across client, sales, and marketing, translating technical constraints into clear requirements and cutting clarification cycles for offshore teams in India and the Philippines.",
        "Shipped and maintained B2B, B2C, and internal dashboards (React, NestJS) plus a React Native app, delivering real-time data views that sharpened decision-making.",
     ]),
    ("Lead Web Developer",
     "FaceCake Marketing Technologies — Los Angeles, CA", "October 2010 – September 2016", [
        "Pioneered early browser-based AR try-on apps (JavaScript, Knockout.js, Pixi.js, computer vision)—the company's first web-AR capability, years before mainstream WebAR.",
        "Built and shipped the NARS Cosmetics virtual try-on app, driving a ~$400K/month revenue lift.",
        "Delivered a Firebase-backed CMS for the NARS team to manage products and content.",
        "Owned the full lifecycle (architecture, development, testing, deployment, maintenance) on a PHP backend.",
     ]),
]

# Projects: (title, subtitle, dates, [ (optional sub-label, text) ... ])
# sub-label = None -> plain bullet; sub-label = str -> bold lead-in delineator
PROJECTS = [
    ("Tidy App — tidyapp.me", "React Native mobile app", "December 2025 – April 2026", [
        (None, "Offline-first, ADHD-friendly home-management app in React Native + Expo + TypeScript."),
        (None, "Prototyped core flows in Lovable, then shipped a production React Native build."),
        (None, "Resilient offline data layer with Zustand and TanStack Query for optimistic UI."),
        (None, "Supabase Auth/PostgREST/Realtime; RevenueCat monetization; Figma Code Connect integration."),
        (None, "90% test coverage with Jest + React Testing Library + Storybook interaction tests."),
    ]),
    ("ForgeKit — forgekit.cloud • npmjs.com", "Open-source Figma-to-React CLI + MCP suite (5,700+ npm installs)", "June 2024 – March 2026", [
        ("Core CLI", "TypeScript CLI scaffolding production-ready Nx monorepos with React 18, Storybook 10+, Vitest, Playwright, and GitHub Actions CI/CD; targets Chakra UI, shadcn/ui, and Tamagui across web and universal React Native; used recursively to scaffold ForgeKit itself."),
        ("Figma MCP", "MCP server extracting Figma variables and design tokens; generates typed theme configs for Chakra UI, Tailwind, and shadcn—enabling AI-driven design-to-code workflows."),
        ("Storybook MCP", "MCP server exposing Storybook metadata, argTypes, and usage patterns to AI coding agents; automates story generation, docs scaffolding, and component-testing workflows."),
    ]),
]

SKILLS = ("React 19, TypeScript, Next.js, Node.js, React Native, Expo, TanStack Start, React Hooks, "
          "Performance Optimization, HTML5, CSS3, Responsive Design, REST & GraphQL Integration, "
          "Storybook, Chromatic, Chakra UI, shadcn/ui, Tamagui, Radix UI, Styled Components / CSS-in-JS, "
          "Model Context Protocol, Claude Code, Cursor, Lovable, Figma Code Connect, Figma MCP, Nx, "
          "Turborepo, Module Federation, GitHub Actions, TanStack Query, Zustand, Supabase, PostgreSQL, "
          "tRPC, Redux Toolkit, Vitest, Jest, Playwright, React Testing Library, a11y, WCAG 2.1 AA, "
          "OpenAI, Agentic Systems")


# ======================= DOCX =======================
def build_docx(path):
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = DOCX_FONT
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.08
    rpr = normal.element.get_or_add_rPr(); rfonts = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rfonts.set(qn(a), DOCX_FONT)
    for s in doc.sections:
        s.top_margin = Inches(0.5); s.bottom_margin = Inches(0.5)
        s.left_margin = Inches(0.7); s.right_margin = Inches(0.7)

    def sp(p, b=0, a=0):
        p.paragraph_format.space_before = Pt(b); p.paragraph_format.space_after = Pt(a)

    def heading(text):
        p = doc.add_paragraph(); sp(p, 8, 2)
        r = p.add_run(text.upper()); r.bold = True; r.font.size = Pt(11.5)
        pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        for k, v in (("w:val", "single"), ("w:sz", "6"), ("w:space", "1"), ("w:color", "000000")):
            bottom.set(qn(k), v)
        pbdr.append(bottom); pPr.append(pbdr)

    def title_line(title):
        p = doc.add_paragraph(); sp(p, 6, 0)
        r = p.add_run(title); r.bold = True; r.font.size = Pt(11)

    def company_line(company, dates):
        p = doc.add_paragraph(); sp(p, 0, 2)
        r = p.add_run(company); r.italic = True; r.font.size = Pt(10)
        p.add_run("\t")
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.1), WD_TAB_ALIGNMENT.RIGHT)
        d = p.add_run(dates); d.font.size = Pt(10)

    def bullet(text, label=None):
        p = doc.add_paragraph(style="List Bullet"); sp(p, 0, 1)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.line_spacing = 1.05
        if label:
            lr = p.add_run(label + " — "); lr.bold = True; lr.font.size = Pt(10.5)
        r = p.add_run(text); r.font.size = Pt(10.5)

    def para(text, a=2):
        p = doc.add_paragraph(); sp(p, 0, a)
        r = p.add_run(text); r.font.size = Pt(10.5)

    # header
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 2)
    r = p.add_run(NAME); r.bold = True; r.font.size = Pt(22)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp(p, 0, 8)
    r = p.add_run(CONTACT); r.font.size = Pt(9.5)

    heading("Summary"); para(SUMMARY)
    heading("Work Experience")
    for title, company, dates, bullets in JOBS:
        title_line(title); company_line(company, dates)
        for b in bullets:
            bullet(b)
    heading("Projects")
    for title, subtitle, dates, items in PROJECTS:
        title_line(title); company_line(subtitle, dates)
        for label, text in items:
            bullet(text, label)
    heading("Skills"); para(SKILLS)

    doc.save(path)


# ======================= PDF =======================
def build_pdf(path):
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                    Spacer, HRFlowable, ListFlowable, ListItem)
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont, TTFError
    from xml.sax.saxutils import escape

    # Prefer Liberation Sans (Arial-metric) for real Unicode glyph extraction; the font
    # directory is configurable and falls back to built-in Helvetica (also Arial-metric)
    # when the TTFs are unavailable (e.g. macOS/Windows or a minimal Linux image).
    fdir = os.environ.get("RESUME_FONT_DIR", "/usr/share/fonts/truetype/liberation/")
    faces = {"ATSSans": "LiberationSans-Regular.ttf", "ATSSans-Bold": "LiberationSans-Bold.ttf",
             "ATSSans-Italic": "LiberationSans-Italic.ttf", "ATSSans-BoldItalic": "LiberationSans-BoldItalic.ttf"}
    try:
        for face, fn in faces.items():
            pdfmetrics.registerFont(TTFont(face, os.path.join(fdir, fn)))
        pdfmetrics.registerFontFamily("ATSSans", normal="ATSSans", bold="ATSSans-Bold",
                                      italic="ATSSans-Italic", boldItalic="ATSSans-BoldItalic")
        FONT, BOLD, ITAL = "ATSSans", "ATSSans-Bold", "ATSSans-Italic"
    except (TTFError, IOError, OSError):
        FONT, BOLD, ITAL = "Helvetica", "Helvetica-Bold", "Helvetica-Oblique"

    doc = BaseDocTemplate(path, pagesize=letter,
                          leftMargin=0.7 * inch, rightMargin=0.7 * inch,
                          topMargin=0.5 * inch, bottomMargin=0.5 * inch,
                          title="Rich Tillman — Resume", author="Rich Tillman")
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="main")
    doc.addPageTemplates([PageTemplate(id="all", frames=[frame])])
    usable = doc.width

    name_st = ParagraphStyle("name", fontName=BOLD, fontSize=20, alignment=TA_CENTER, spaceAfter=2, leading=22)
    contact_st = ParagraphStyle("contact", fontName=FONT, fontSize=9, alignment=TA_CENTER, spaceAfter=6, leading=11)
    head_st = ParagraphStyle("head", fontName=BOLD, fontSize=11, spaceBefore=7, spaceAfter=2, leading=12)
    title_st = ParagraphStyle("title", fontName=BOLD, fontSize=10.5, spaceBefore=5, spaceAfter=0, leading=12)
    meta_st = ParagraphStyle("meta", fontName=FONT, fontSize=9.5, leading=11, spaceAfter=2)
    body_st = ParagraphStyle("body", fontName=FONT, fontSize=10, leading=12, spaceAfter=1)
    bullet_st = ParagraphStyle("bul", fontName=FONT, fontSize=10, leading=12, spaceAfter=1.5)

    story = []

    def heading(text):
        story.append(Paragraph(text.upper(), head_st))
        story.append(HRFlowable(width="100%", thickness=0.6, color=colors.black,
                                spaceBefore=1, spaceAfter=3, lineCap="square"))

    def title_company(title, company, dates):
        # Single-line, natural reading order (no layout tables) — company italic, dates inline.
        story.append(Paragraph(escape(title), title_st))
        story.append(Paragraph("<i>%s</i>&nbsp;&nbsp;—&nbsp;&nbsp;%s"
                               % (escape(company), escape(dates)), meta_st))

    def bullets(items):
        li = []
        for it in items:
            if isinstance(it, tuple):
                label, text = it
                html = ("<b>%s — </b>%s" % (escape(label), escape(text))) if label else escape(text)
            else:
                html = escape(it)
            li.append(ListItem(Paragraph(html, bullet_st), value=None, leftIndent=16))
        story.append(ListFlowable(li, bulletType="bullet", bulletFontName=FONT, bulletFontSize=8,
                                  bulletColor=colors.black, leftIndent=14, start="•",
                                  spaceBefore=0, spaceAfter=0))

    story.append(Paragraph(escape(NAME), name_st))
    story.append(Paragraph(escape(CONTACT), contact_st))
    heading("Summary"); story.append(Paragraph(escape(SUMMARY), body_st))
    heading("Work Experience")
    for title, company, dates, bl in JOBS:
        title_company(title, company, dates); bullets(bl)
    heading("Projects")
    for title, subtitle, dates, items in PROJECTS:
        title_company(title, subtitle, dates); bullets(items)
    heading("Skills"); story.append(Paragraph(escape(SKILLS), body_st))

    doc.build(story)


if __name__ == "__main__":
    os.makedirs(OUTDIR, exist_ok=True)
    docx_path = os.path.join(OUTDIR, BASENAME + ".docx")
    pdf_path = os.path.join(OUTDIR, BASENAME + ".pdf")
    build_docx(docx_path); print("Saved", docx_path)
    build_pdf(pdf_path); print("Saved", pdf_path)
